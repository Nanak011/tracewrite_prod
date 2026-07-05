from __future__ import annotations

import argparse
import base64
import json
import re
import sys
from dataclasses import dataclass, field
from enum import Enum, auto
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.document import Document as _Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# ---------------------------------------------------------------------------
# Optional NLP backends - imported lazily, never required for tiers 1 & 2
# ---------------------------------------------------------------------------

_NLTK_READY = False
_SPACY_NLP = None  # lazily initialized singleton, may stay None forever


def _ensure_nltk() -> bool:
    """Best-effort NLTK setup. Returns True if POS tagging is usable."""
    global _NLTK_READY
    if _NLTK_READY:
        return True
    try:
        import nltk
        for resource, sub in (
            ("punkt", "tokenizers/punkt"),
            ("averaged_perceptron_tagger", "taggers/averaged_perceptron_tagger"),
            ("averaged_perceptron_tagger_eng", "taggers/averaged_perceptron_tagger_eng"),
        ):
            try:
                nltk.data.find(sub)
            except LookupError:
                try:
                    nltk.download(resource, quiet=True)
                except Exception:
                    pass
        _NLTK_READY = True
        return True
    except Exception:
        return False


def _ensure_spacy():
    """Best-effort spaCy setup. Returns a pipeline or None. Never raises."""
    global _SPACY_NLP
    if _SPACY_NLP is not None:
        return _SPACY_NLP
    try:
        import spacy
        try:
            _SPACY_NLP = spacy.load("en_core_web_sm")
        except Exception:
            _SPACY_NLP = spacy.blank("en")
    except Exception:
        _SPACY_NLP = False  # sentinel: spaCy not installed at all
    return _SPACY_NLP or None


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

# Tier 1: "1", "1.", "1.1", "1.1.1" followed by a title. Max depth 3.
# Accepts every spacing/punctuation variant around the number:
# "1 heading", "1heading", "1Heading", "1.Heading", "1. Heading", "1heading.",
# "1 Heading.", "1. Heading.", "1. heading." - the period after the number
# and the space before the title are both optional now, not required.
_HEADING_RE = re.compile(r"^(\d{1,2}(?:\.\d{1,2}){0,2})(?!\))\.?\s*(\S.*)$")

# Short, single-word "titles" that are really just units/times stuck to a
# number with no separator ("3pm", "5ft", "10km") - without this, loosening
# the regex above to allow no-space numbering would misclassify these as
# headings. Only blocks single-word matches; "10km race day" still passes.
_NUMBER_SUFFIX_GUARD = {
    "am", "pm", "kg", "km", "cm", "mm", "ft", "in", "lb", "lbs",
    "st", "nd", "rd", "th", "hz", "khz", "mhz", "ghz", "mph",
    "kmh", "gb", "mb", "kb", "pt", "px", "fps", "v", "w", "k", "x",
}

# Tier 2: list markers that do NOT collide with the heading convention.
# Plain "N." is intentionally excluded - it is reserved for headings.
_BULLET_CHARS = {"-", "*", "\u2022", "\u2023", "\u25E6", "\u2043", "\u2219"}
_LIST_PATTERNS = [
    re.compile(r"^\d{1,3}\)"),        # 1)
    re.compile(r"^\(\d{1,3}\)"),      # (1)
    re.compile(r"^[a-z]\."),          # a.
    re.compile(r"^[a-z]\)"),          # a)
    re.compile(r"^\([a-z]\)"),        # (a)
    re.compile(r"^[ivxlcdm]+\."),     # iv.
    re.compile(r"^\([ivxlcdm]+\)"),   # (iv)
]

_BLOCKQUOTE_RE = re.compile(r"^>\s?")
_CODE_FIRST_WORDS = {
    "def", "class", "import", "from", "return", "if", "elif", "else",
    "for", "while", "try", "except", "with", "async", "await", "print",
    "const", "let", "var", "function", "public", "private", "protected",
    "package", "include", "fn", "impl", "use", "mod", "pub", "struct",
    "enum", "trait", "match",
}


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

class BlockType(Enum):
    PARAGRAPH = auto()
    HEADING = auto()
    LIST_ITEM_ORDERED = auto()
    LIST_ITEM_UNORDERED = auto()
    BLOCKQUOTE = auto()
    CODE_BLOCK = auto()
    TITLE = auto()
    ABSTRACT = auto()
    PAGE_BREAK = auto()
    EMPTY = auto()


@dataclass
class RunFormat:
    text: str = ""
    bold: bool = False
    italic: bool = False
    underline: bool = False
    font_name: Optional[str] = None
    font_size: Optional[float] = None


@dataclass
class HyperlinkInfo:
    text: str = ""
    url: str = ""


@dataclass
class ExtractedBlock:
    type: BlockType = BlockType.PARAGRAPH
    text: str = ""
    runs: List[RunFormat] = field(default_factory=list)
    hyperlinks: List[HyperlinkInfo] = field(default_factory=list)
    heading_level: int = 0
    heading_number: Optional[str] = None
    heading_flagged: bool = False  # True = NLP fallback guess, needs confirmation
    list_item_marker: Optional[str] = None
    list_level: int = 0
    left_indent: Optional[float] = None  # inches, from the source doc - used to detect nested list depth
    original_index: int = 0


@dataclass
class StyleConfig:
    """User-tunable formatting knobs, surfaced 1:1 from the frontend panel."""
    font_name: str = "Times New Roman"
    heading_font_name: Optional[str] = None
    monospace_font_name: str = "Consolas"
    body_font_size_pt: int = 12
    heading_1_size_pt: int = 16
    heading_2_size_pt: int = 14
    heading_3_size_pt: int = 13
    heading_1_bold: bool = True
    heading_2_bold: bool = True
    heading_3_bold: bool = True
    heading_1_color: Optional[str] = None
    heading_2_color: Optional[str] = None
    heading_3_color: Optional[str] = None
    line_spacing: float = 1.5
    paragraph_space_before_pt: int = 0
    paragraph_space_after_pt: int = 6
    heading_1_space_before_pt: int = 18
    heading_1_space_after_pt: int = 10
    heading_2_space_before_pt: int = 12
    heading_2_space_after_pt: int = 6
    heading_3_space_before_pt: int = 8
    heading_3_space_after_pt: int = 4
    page_margin_top_in: float = 1.0
    page_margin_bottom_in: float = 1.0
    page_margin_left_in: float = 1.0
    page_margin_right_in: float = 1.0
    body_alignment: str = "justify"
    heading_alignment: str = "left"
    title_alignment: str = "center"
    first_line_indent_in: float = 0.3
    enable_first_line_indent: bool = True
    blockquote_indent_in: float = 0.5
    list_indent_in: float = 0.25
    list_bullet_char: str = "\u2022"
    include_toc: bool = True
    toc_title: str = "Table of Contents"
    toc_max_level: int = 3
    # NLP fallback tier
    enable_nlp_fallback: bool = True
    nlp_min_heading_score: float = 0.6
    detect_blockquotes: bool = True
    detect_code_blocks: bool = True
    detect_abstract: bool = True

    def __post_init__(self):
        if self.heading_font_name is None:
            self.heading_font_name = self.font_name

    @classmethod
    def from_dict(cls, d: dict) -> "StyleConfig":
        allowed = set(StyleConfig.__dataclass_fields__.keys())
        return cls(**{k: v for k, v in d.items() if k in allowed})


# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------

def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _load_style_config(path: Path) -> StyleConfig:
    if not path.exists():
        return StyleConfig()
    raw = json.loads(path.read_text(encoding="utf-8"))
    return StyleConfig.from_dict(raw)


def _para_alignment(value: str):
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(value.lower(), WD_ALIGN_PARAGRAPH.JUSTIFY)


def _rgb(color_hex: Optional[str]) -> Optional[RGBColor]:
    if not color_hex:
        return None
    h = color_hex.lstrip("#")
    try:
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    except (ValueError, IndexError):
        return None


def _iter_paragraphs(doc: _Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        # tables intentionally skipped - out of scope for this version


def _set_paragraph_spacing(paragraph: Paragraph, before=None, after=None,
                            line_spacing=None) -> None:
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing


# ---------------------------------------------------------------------------
# Extraction (DOCX only - this pipeline's input is always Quill -> DOCX)
# ---------------------------------------------------------------------------

def _extract_run_format(run: Run) -> RunFormat:
    f = run.font
    return RunFormat(
        text=run.text,
        bold=bool(run.bold),
        italic=bool(run.italic),
        underline=bool(run.underline),
        font_name=f.name,
        font_size=f.size.pt if f.size else None,
    )


def _extract_hyperlinks(paragraph: Paragraph) -> List[HyperlinkInfo]:
    results: List[HyperlinkInfo] = []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for hl in paragraph._p.findall(".//w:hyperlink", ns):
        rel_id = hl.get(qn("r:id"))
        url = ""
        if rel_id:
            rel = paragraph.part.rels.get(rel_id)
            if rel:
                url = rel.target_ref or ""
        texts = [t.text for r in hl.findall(".//w:r", ns)
                 for t in [r.find("w:t", ns)] if t is not None and t.text]
        results.append(HyperlinkInfo(text="".join(texts), url=url))
    return results


def extract_blocks(path: Path) -> List[ExtractedBlock]:
    doc = Document(str(path))
    blocks: List[ExtractedBlock] = []
    idx = 0
    for para in _iter_paragraphs(doc):
        text = _clean(para.text)
        if not text:
            is_page_break = any(
                br.get(qn("w:type")) == "page"
                for br in para._p.findall(qn("w:br"))
            )
            if is_page_break:
                blocks.append(ExtractedBlock(type=BlockType.PAGE_BREAK, original_index=idx))
                idx += 1
            continue
        runs = [_extract_run_format(r) for r in para.runs if r.text.strip()]
        hyperlinks = _extract_hyperlinks(para)
        pf = para.paragraph_format
        left_indent = pf.left_indent.inches if pf.left_indent else None
        blocks.append(ExtractedBlock(
            type=BlockType.PARAGRAPH, text=text, runs=runs,
            hyperlinks=hyperlinks, left_indent=left_indent, original_index=idx,
        ))
        idx += 1
    return blocks


# ---------------------------------------------------------------------------
# Tier 3: NLP fallback classifier
# ---------------------------------------------------------------------------

class NlpFallbackClassifier:
    """
    Scores a short, period-free, capitalized line for "does this look
    like a heading" using POS structure and English grammar heuristics
    rather than formatting cues. Only ever called on lines that already
    survived the cheap heuristic gate in HeadingPipeline._heuristic_prefilter,
    but the no-fullstop check is repeated here too, so this scorer gives a
    sane answer even if called directly elsewhere.
    """

    # Common section-heading vocabulary across academic/report writing.
    # A line built mostly from these words is heading-shaped almost by
    # definition, independent of POS structure.
    _HEADING_KEYWORDS = {
        "introduction", "background", "overview", "summary", "abstract",
        "conclusion", "conclusions", "discussion", "methodology", "methods",
        "results", "findings", "analysis", "evaluation", "experiment",
        "experiments", "setup", "design", "approach", "implementation",
        "related", "work", "literature", "review", "motivation", "problem",
        "objectives", "scope", "limitations", "future", "recommendations",
        "appendix", "references", "acknowledgments", "acknowledgements",
        "glossary", "definitions", "terminology", "requirements",
    }

    def __init__(self):
        self._nltk_ok = _ensure_nltk()
        self._spacy = _ensure_spacy() if _ensure_nltk else None

    def score(self, text: str) -> float:
        words = text.split()
        if not words:
            return 0.0
        # No-fullstop / phrase rule: a heading is a label, not a sentence.
        # This is already enforced upstream (_heuristic_prefilter), but
        # repeated here so score() is safe to call on its own too.
        if text.rstrip().endswith((".", "!", "?")):
            return 0.0

        # Title-case ratio
        title_ratio = sum(1 for w in words if w[:1].isupper()) / len(words)
        tc_score = 0.35 if title_ratio >= 0.55 else (0.15 if title_ratio >= 0.3 else 0.0)

        # Common heading vocabulary - matched independent of case/punctuation
        bare_words = {w.strip(".,;:!?").lower() for w in words}
        keyword_hits = len(bare_words & self._HEADING_KEYWORDS)
        keyword_bonus = min(0.3, keyword_hits * 0.2)

        # POS: noun-heavy / no finite-verb-as-clause signal
        noun_score = 0.0
        verb_penalty = 0.0
        opener_penalty = 0.0
        if self._nltk_ok:
            try:
                import nltk
                tagged = nltk.pos_tag(words)
                noun_ratio = sum(1 for _, t in tagged if t.startswith("NN")) / len(tagged)
                noun_score = 0.3 if noun_ratio >= 0.35 else (0.15 if noun_ratio >= 0.2 else 0.0)
                # Finite verbs/modals (VBZ/VBD/VBP/MD) mean this is a full
                # clause ("This shows...", "We will..."), not a heading
                # label. VBG/VBN are excluded deliberately - gerund/participle
                # openers like "Understanding X" or "Related Work" are
                # completely normal, common heading phrasing.
                if any(t in ("VBZ", "VBD", "VBP", "MD") for _, t in tagged):
                    verb_penalty = 0.25
                # Headings rarely open with a pronoun or conjunction
                # ("It...", "This...", "And...") - that's a sentence opener.
                if tagged and tagged[0][1] in ("PRP", "CC", "IN"):
                    opener_penalty = 0.15
            except Exception:
                pass

        # Optional spaCy entity-density bonus (purely additive, never required)
        ent_bonus = 0.0
        if self._spacy:
            try:
                doc = self._spacy(text)
                ent_bonus = min(0.15, len(list(doc.ents)) * 0.05)
            except Exception:
                pass

        colon_bonus = 0.1 if ":" in text and len(text.split(":")[0].split()) <= 5 else 0.0
        score = (tc_score + noun_score + keyword_bonus + ent_bonus + colon_bonus
                 - verb_penalty - opener_penalty)
        return max(0.0, min(1.0, score))


# ---------------------------------------------------------------------------
# Heading + list + blockquote + code detection pipeline
# ---------------------------------------------------------------------------

class HeadingPipeline:
    def __init__(self, cfg: StyleConfig):
        self.cfg = cfg
        self._nlp = NlpFallbackClassifier() if cfg.enable_nlp_fallback else None
        self.warnings: List[str] = []
        self._counters = [0, 0, 0]
        self._last_confirmed_level = 0
        self._last_confirmed_number: Optional[str] = None
        self._used_numbers: Dict[str, int] = {}

    # -- Tier 1 ------------------------------------------------------------
    @staticmethod
    def _match_numbered(text: str) -> Optional[Tuple[List[int], str]]:
        m = _HEADING_RE.match(text)
        if not m:
            return None
        number_str, title = m.group(1), _clean(m.group(2))
        parts = [int(p) for p in number_str.split(".")]
        if len(parts) > 3 or not title:
            return None
        # Reject obvious sentences that happen to start with a number,
        # e.g. "2024 was a good year for research." - long + ends in period
        if title.endswith((".", ",", ";")) and len(title.split()) > 6:
            return None
        if len(title.split()) > 15:
            return None
        # Guard against no-space numbering swallowing units/times, e.g.
        # "3pm", "5ft", "10km" - these would otherwise match as a heading
        # numbered "3"/"5"/"10" titled "pm"/"ft"/"km".
        stripped_title = title.rstrip(".").strip()
        if len(stripped_title.split()) == 1 and stripped_title.lower() in _NUMBER_SUFFIX_GUARD:
            return None
        # A short trailing period is just punctuation, not a real sentence
        # end - strip it so "1. Heading." renders as "Heading", not
        # "Heading." (the long-sentence case above is already rejected).
        title = stripped_title
        return parts, title

    # -- Tier 2 --------------------------------------------------------------
    @staticmethod
    def match_list_item(text: str) -> Optional[str]:
        t = _clean(text)
        if not t:
            return None
        if t[0] in _BULLET_CHARS and t[1:].strip():
            return t[0]
        for pat in _LIST_PATTERNS:
            m = pat.match(t)
            if m and t[len(m.group(0)):].strip():
                return m.group(0)
        return None

    # -- Tier 3 gate ---------------------------------------------------------
    @staticmethod
    def _heuristic_prefilter(text: str) -> bool:
        words = text.split()
        if not (1 <= len(words) <= 12):
            return False
        if text.endswith((".", ",", ";")):
            return False
        if not words[0][:1].isupper():
            return False
        return True

    def _assign_number(self, level: int, line: int, text: str,
                        explicit_parts: Optional[List[int]] = None) -> str:
        """
        Advances the real counters and returns a number string. Used by
        BOTH tiers: tier 1 passes explicit_parts (the author's own
        digits), tier 3 passes None (advance one level under the last
        confirmed heading ourselves).

        If the resulting number was already used earlier in the document
        (e.g. an NLP-guessed heading claimed "2.1", then the author's own
        next heading is also literally "2.1" because they lost count),
        this bumps the LAST segment forward until it finds a free number
        - so the final document never shows a literal duplicate - and
        keeps the internal counters in sync with whatever number actually
        got used, so anything numbered relative to it afterward stays
        consistent. The change is always reported in the warnings; we
        never silently renumber something the author explicitly typed.
        """
        if explicit_parts:
            for i in range(3):
                self._counters[i] = explicit_parts[i] if i < len(explicit_parts) else 0
            level = len(explicit_parts)
        else:
            self._counters[level - 1] += 1
            for i in range(level, 3):
                self._counters[i] = 0

        def _format(counters: List[int]) -> str:
            return ".".join(str(c) for c in counters[:level])

        original = _format(self._counters)
        number = original
        while number in self._used_numbers:
            self._counters[level - 1] += 1
            number = _format(self._counters)

        if number != original:
            self.warnings.append(
                f"Line {line}: heading \u201c{text}\u201d was numbered {original}, which "
                f"conflicts with an earlier heading. It was automatically renumbered "
                f"to {number} - please check this is the right place for it."
            )

        self._used_numbers[number] = line
        self._last_confirmed_level = level
        self._last_confirmed_number = number
        return number

    def run(self, blocks: List[ExtractedBlock]) -> None:
        for i, block in enumerate(blocks):
            if block.type != BlockType.PARAGRAPH:
                continue
            text = _clean(block.text)
            if not text:
                continue

            # Tier 1
            numbered = self._match_numbered(text)
            if numbered:
                parts, title = numbered
                level = len(parts)
                number = self._assign_number(level, block.original_index, title, explicit_parts=parts)
                block.type = BlockType.HEADING
                block.heading_level = level
                block.heading_number = number
                block.heading_flagged = False
                block.text = title
                continue

            # Tier 2 (checked before tier 3 so list items never get
            # miscast as fallback headings)
            marker = self.match_list_item(text)
            if marker:
                is_ordered = marker not in _BULLET_CHARS
                item_text = text[len(marker):].strip() if is_ordered else text[1:].strip()
                block.type = (BlockType.LIST_ITEM_ORDERED if is_ordered
                              else BlockType.LIST_ITEM_UNORDERED)
                block.list_item_marker = marker if is_ordered else None
                block.text = item_text
                continue

            # Tier 3
            if self._nlp and self._heuristic_prefilter(text):
                score = self._nlp.score(text)
                if score >= self.cfg.nlp_min_heading_score:
                    level = min(self._last_confirmed_level + 1, 3) or 1
                    number = self._assign_number(level, block.original_index, text)
                    block.type = BlockType.HEADING
                    block.heading_level = level
                    block.heading_number = number
                    block.heading_flagged = True
                    self.warnings.append(
                        f"Line {block.original_index}: \u201c{text}\u201d looks like a "
                        f"heading but had no number, so it was given \u201c{number}\u201d "
                        f"automatically. Please double-check this is the right number."
                    )
                    continue

            # Neither tier matched -> leave as body text. Do not guess further.


def _detect_blockquotes(blocks: List[ExtractedBlock]) -> None:
    for b in blocks:
        if b.type != BlockType.PARAGRAPH:
            continue
        text = _clean(b.text)
        if _BLOCKQUOTE_RE.match(text):
            b.type = BlockType.BLOCKQUOTE
            b.text = _BLOCKQUOTE_RE.sub("", text).strip()


def _detect_code_blocks(blocks: List[ExtractedBlock]) -> None:
    for b in blocks:
        if b.type != BlockType.PARAGRAPH:
            continue
        text = _clean(b.text)
        if not text:
            continue
        if b.runs:
            mono = sum(1 for r in b.runs if r.font_name and
                       any(k in r.font_name.lower() for k in
                           ("mono", "courier", "consolas", "menlo")) and r.text.strip())
            total = sum(1 for r in b.runs if r.text.strip())
            if total > 0 and mono / total >= 0.6:
                b.type = BlockType.CODE_BLOCK
                continue
        first_word = text.split()[0] if text.split() else ""
        if first_word in _CODE_FIRST_WORDS:
            b.type = BlockType.CODE_BLOCK


def _detect_list_levels(blocks: List[ExtractedBlock]) -> None:
    """
    Word/Quill nested lists are visually indented in the source doc, but
    that indentation is just a paragraph_format.left_indent value - it
    doesn't carry an explicit "level 2" marker anywhere. We recover
    nesting depth by clustering consecutive list items and mapping their
    distinct left_indent values to levels 0, 1, 2... within that cluster.
    Each cluster is scoped independently so an unrelated list later in
    the document doesn't shift this list's levels.
    """
    i = 0
    while i < len(blocks):
        if blocks[i].type not in (BlockType.LIST_ITEM_ORDERED, BlockType.LIST_ITEM_UNORDERED):
            i += 1
            continue
        j = i
        while j < len(blocks) and blocks[j].type in (
                BlockType.LIST_ITEM_ORDERED, BlockType.LIST_ITEM_UNORDERED):
            j += 1
        cluster = blocks[i:j]
        indents = sorted({(b.left_indent or 0.0) for b in cluster})
        level_map = {v: lvl for lvl, v in enumerate(indents)}
        for b in cluster:
            b.list_level = level_map[b.left_indent or 0.0]
        i = j


def _detect_abstract(blocks: List[ExtractedBlock]) -> None:
    for i, b in enumerate(blocks):
        if b.type != BlockType.PARAGRAPH:
            continue
        text = _clean(b.text).lower()
        if text in ("abstract", "summary") and i + 1 < len(blocks):
            if blocks[i + 1].type == BlockType.PARAGRAPH:
                blocks[i + 1].type = BlockType.ABSTRACT
            b.type = BlockType.EMPTY
        break  # only ever check the very first candidate block


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def _apply_heading_style(paragraph: Paragraph, level: int, cfg: StyleConfig,
                          number: Optional[str], title: str) -> None:
    """
    Applies a REAL Word heading style (Heading 1/2/3), not just a bold
    Normal paragraph. This matters: Word's navigation pane, outline view,
    screen readers, and any native TOC field the user inserts later all
    key off the paragraph style, not the run's bold flag. cfg-driven
    font/size/color/spacing are then layered on top at the run level so
    the visual result still matches the user's chosen settings exactly.

    `number` is None for NLP-fallback-flagged headings - those render as
    a normal, clean heading with no number, since we don't fabricate a
    number the author didn't write. The heading is otherwise identical
    to a confirmed one; there is no italic/asterisk/"[unconfirmed]" text
    in the document. The one-time review prompt lives entirely in the
    app's warnings panel, not in the downloadable file.
    """
    sizes = {1: cfg.heading_1_size_pt, 2: cfg.heading_2_size_pt, 3: cfg.heading_3_size_pt}
    bolds = {1: cfg.heading_1_bold, 2: cfg.heading_2_bold, 3: cfg.heading_3_bold}
    colors = {1: cfg.heading_1_color, 2: cfg.heading_2_color, 3: cfg.heading_3_color}
    sp_before = {1: cfg.heading_1_space_before_pt, 2: cfg.heading_2_space_before_pt,
                 3: cfg.heading_3_space_before_pt}
    sp_after = {1: cfg.heading_1_space_after_pt, 2: cfg.heading_2_space_after_pt,
                3: cfg.heading_3_space_after_pt}

    font_name = cfg.heading_font_name or cfg.font_name
    full_text = f"{number} {title}" if number else title

    style_name = f"Heading {min(max(level, 1), 3)}"
    try:
        paragraph.style = paragraph.part.document.styles[style_name]
    except KeyError:
        pass  # falls back to whatever style add_paragraph() gave it

    # Clear whatever placeholder text the built-in style run carries and
    # write our own run so every property below is explicit, not inherited.
    run = paragraph.add_run(full_text)
    run.font.name = font_name
    run.font.size = Pt(sizes.get(level, cfg.body_font_size_pt))
    run.bold = bolds.get(level, True)
    # IMPORTANT: Word's built-in Heading 1/2/3 styles carry their own blue
    # theme color. If we don't set an explicit color here, the run just
    # inherits that blue instead of respecting "no color chosen = black".
    # So this is never conditional - always set something.
    rgb = _rgb(colors.get(level)) or RGBColor(0, 0, 0)
    run.font.color.rgb = rgb

    paragraph.alignment = _para_alignment(cfg.heading_alignment)
    pf = paragraph.paragraph_format
    pf.space_before = Pt(sp_before.get(level, cfg.paragraph_space_before_pt))
    pf.space_after = Pt(sp_after.get(level, cfg.paragraph_space_after_pt))
    pf.first_line_indent = None  # headings never get the body first-line indent
    pf.left_indent = None
    pf.keep_with_next = True


def _reconstruct_runs(paragraph: Paragraph, runs: List[RunFormat], cfg: StyleConfig) -> None:
    """Preserves per-run bold/italic/font. Only safe to use when block.text
    was NOT modified after extraction (i.e. runs and text still agree)."""
    if not runs:
        return
    for rf in runs:
        if not rf.text.strip():
            continue
        run = paragraph.add_run(rf.text)
        run.font.name = rf.font_name or cfg.font_name
        run.font.size = Pt(rf.font_size or cfg.body_font_size_pt)
        run.bold = rf.bold
        run.italic = rf.italic
        run.underline = rf.underline


def _dominant_format(runs: List[RunFormat]) -> RunFormat:
    """First non-empty run's formatting, used as a single style for a whole
    line when block.text has been rewritten and per-run character offsets
    (list markers stripped, blockquote '>' stripped) can no longer be trusted."""
    for r in runs:
        if r.text.strip():
            return r
    return RunFormat()


def _build_toc(doc: Document, blocks: List[ExtractedBlock], cfg: StyleConfig) -> None:
    """
    Inserts a REAL Word TOC field (TOC \\o "1-N" \\h \\z \\u) using the
    standard single-paragraph field structure: begin -> instrText ->
    separate -> placeholder text -> end, all within one paragraph. This
    matches exactly what Word itself generates when you insert a TOC via
    the ribbon.

    An earlier version tried to pre-populate the field's cached region
    with our own multi-paragraph heading list (so something readable
    showed before the user ever updated the field). That turned out to
    cause every TOC entry to resolve to "page 1" when Word recomputed it
    - real Word-generated TOCs carry internal hyperlink/bookmark
    sub-structure per entry that a hand-rolled multi-paragraph cache
    doesn't have, and Word's update algorithm apparently falls back to a
    default rather than erroring visibly. Correctness (real page numbers)
    matters more than a pre-filled placeholder, so this reverts to the
    simple, well-tested structure - the trade-off is the user sees
    placeholder text until they update the field (or Word does it
    automatically via w:updateFields, see _force_update_fields_on_open).
    """
    if not cfg.include_toc:
        return
    toc_heading = doc.add_heading(cfg.toc_title, level=1)
    toc_color = _rgb(cfg.heading_1_color) or RGBColor(0, 0, 0)
    for run in toc_heading.runs:
        run.font.color.rgb = toc_color
        run.font.name = cfg.heading_font_name or cfg.font_name

    has_headings = any(b.type == BlockType.HEADING for b in blocks)
    paragraph = doc.add_paragraph()

    if not has_headings:
        paragraph.add_run("(No headings found)")
        doc.add_page_break()
        return

    run = paragraph.add_run()
    r = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-%d" \\h \\z \\u' % cfg.toc_max_level
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.set(qn("xml:space"), "preserve")
    placeholder.text = "Right-click here and choose \u201cUpdate Field\u201d to build the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)
    r.append(fld_end)

    doc.add_page_break()


def _add_page_numbers(doc: Document) -> None:
    """Adds a centered PAGE field to the footer. One section covers the
    whole document here, so this applies to every page automatically -
    front matter included."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_end)


def _force_update_fields_on_open(doc: Document) -> None:
    """
    Tells Word to recompute all fields - the TOC field and the page
    number field - the moment the file opens, using Word's own layout
    engine. This is what turns the pre-populated TOC (real headings, no
    page numbers - see _build_toc) into a genuinely real TOC with actual
    page numbers and dot leaders, automatically, without the user
    right-clicking "Update Field" themselves.

    Trade-off: some Word configurations show a one-time "This document
    contains fields... update them?" prompt when this is set. We're
    choosing correctness (a real TOC) over avoiding that prompt - the
    cached content from _build_toc is still there as a reasonable
    fallback for viewers that ignore this setting entirely (e.g. some
    LibreOffice/Google Docs import paths).
    """
    settings_el = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings_el.append(update_fields)


def build_document(blocks: List[ExtractedBlock], output_path: Path, cfg: StyleConfig) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(cfg.page_margin_top_in)
    sec.bottom_margin = Inches(cfg.page_margin_bottom_in)
    sec.left_margin = Inches(cfg.page_margin_left_in)
    sec.right_margin = Inches(cfg.page_margin_right_in)

    _build_toc(doc, blocks, cfg)
    _add_page_numbers(doc)

    just_started_section = True  # first paragraph after a heading is never indented

    for block in blocks:
        if block.type == BlockType.EMPTY:
            continue
        if block.type == BlockType.PAGE_BREAK:
            doc.add_page_break()
            continue

        if block.type == BlockType.HEADING:
            p = doc.add_paragraph()
            _apply_heading_style(p, block.heading_level, cfg,
                                  block.heading_number, block.text)
            just_started_section = True
            continue

        if block.type in (BlockType.LIST_ITEM_ORDERED, BlockType.LIST_ITEM_UNORDERED):
            p = doc.add_paragraph(style="List Paragraph")
            fmt = _dominant_format(block.runs)
            is_ordered = block.type == BlockType.LIST_ITEM_ORDERED
            prefix = f"{block.list_item_marker} " if is_ordered else f"{cfg.list_bullet_char} "
            run = p.add_run(prefix + block.text)
            run.font.name = fmt.font_name or cfg.font_name
            run.font.size = Pt(fmt.font_size or cfg.body_font_size_pt)
            run.bold = fmt.bold
            run.italic = fmt.italic
            # Numbered lists sit one indent-unit in; bullets sit two units
            # in, so the two marker types are visually distinct even at
            # the same nesting depth. Deeper nesting (detected from the
            # source doc's own indentation) adds further indent on top.
            base_units = 1 if is_ordered else 2
            indent = cfg.list_indent_in * (base_units + block.list_level)
            p.paragraph_format.left_indent = Inches(indent)
            p.paragraph_format.first_line_indent = Inches(-cfg.list_indent_in)
            _set_paragraph_spacing(p, after=3, line_spacing=cfg.line_spacing)
            just_started_section = True
            continue

        if block.type == BlockType.BLOCKQUOTE:
            p = doc.add_paragraph()
            fmt = _dominant_format(block.runs)
            run = p.add_run(block.text)
            run.font.name = fmt.font_name or cfg.font_name
            run.font.size = Pt(fmt.font_size or cfg.body_font_size_pt)
            run.italic = True
            p.paragraph_format.left_indent = Inches(cfg.blockquote_indent_in)
            _set_paragraph_spacing(p, after=cfg.paragraph_space_after_pt, line_spacing=cfg.line_spacing)
            just_started_section = True
            continue

        if block.type == BlockType.CODE_BLOCK:
            p = doc.add_paragraph()
            run = p.add_run(block.text)
            run.font.name = cfg.monospace_font_name
            run.font.size = Pt(max(cfg.body_font_size_pt - 1, 8))
            p.paragraph_format.left_indent = Inches(0.5)
            _set_paragraph_spacing(p, before=6, after=6)
            just_started_section = True
            continue

        if block.type == BlockType.ABSTRACT:
            doc.add_paragraph().add_run("Abstract: ").italic = True
            p = doc.add_paragraph()
            _reconstruct_runs(p, block.runs, cfg)
            p.alignment = _para_alignment("justify")
            _set_paragraph_spacing(p, after=cfg.paragraph_space_after_pt, line_spacing=cfg.line_spacing)
            doc.add_page_break()  # body starts fresh after front matter
            just_started_section = True
            continue

        # Default: body paragraph. First paragraph of a section (right
        # after a heading) renders flush left with no first-line indent -
        # standard typographic convention. Only paragraphs CONTINUING a
        # section get the indent.
        p = doc.add_paragraph()
        _reconstruct_runs(p, block.runs, cfg)
        p.alignment = _para_alignment(cfg.body_alignment)
        _set_paragraph_spacing(p, before=cfg.paragraph_space_before_pt,
                                after=cfg.paragraph_space_after_pt,
                                line_spacing=cfg.line_spacing)
        if (cfg.enable_first_line_indent and cfg.first_line_indent_in > 0
                and not just_started_section):
            p.paragraph_format.first_line_indent = Inches(cfg.first_line_indent_in)
        just_started_section = False

    _force_update_fields_on_open(doc)
    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------

def _html_escape(text: str) -> str:
    return (text.replace("&", "&amp;").replace("<", "&lt;")
                .replace(">", "&gt;").replace('"', "&quot;"))


def build_preview_html(blocks: List[ExtractedBlock], cfg: StyleConfig) -> str:
    """
    Renders the SAME blocks+cfg used for the real .docx into inline-styled
    HTML for the in-browser preview. This is deliberately NOT a re-parse
    of the generated docx (e.g. via mammoth) - round-tripping through a
    second converter is exactly how the preview drifted from the real
    file before (mammoth drops direct/inline formatting like first-line
    indent and exact heading colors, keeping only named paragraph styles).
    Generating both outputs from one source means they can't drift apart.
    """
    heading_font = cfg.heading_font_name or cfg.font_name
    sizes = {1: cfg.heading_1_size_pt, 2: cfg.heading_2_size_pt, 3: cfg.heading_3_size_pt}
    colors = {1: cfg.heading_1_color, 2: cfg.heading_2_color, 3: cfg.heading_3_color}

    parts: List[str] = []
    just_started_section = True

    for block in blocks:
        if block.type == BlockType.EMPTY:
            continue
        if block.type == BlockType.PAGE_BREAK:
            parts.append('<hr style="border-top: 1px dashed #ccc; margin: 2em 0;">')
            continue

        if block.type == BlockType.HEADING:
            lvl = min(max(block.heading_level, 1), 3)
            color = f"#{colors.get(lvl)}" if colors.get(lvl) else "#000000"
            label = f"{block.heading_number} {block.text}" if block.heading_number else block.text
            parts.append(
                f'<h{lvl} style="font-family:\'{heading_font}\',serif; '
                f'font-size:{sizes.get(lvl, cfg.body_font_size_pt)}pt; color:{color}; '
                f'font-weight:bold; margin:1em 0 0.3em;">{_html_escape(label)}</h{lvl}>'
            )
            just_started_section = True
            continue

        if block.type in (BlockType.LIST_ITEM_ORDERED, BlockType.LIST_ITEM_UNORDERED):
            is_ordered = block.type == BlockType.LIST_ITEM_ORDERED
            prefix = f"{block.list_item_marker} " if is_ordered else f"{cfg.list_bullet_char} "
            base_units = 1 if is_ordered else 2
            indent_em = 1.5 * (base_units + block.list_level)
            parts.append(
                f'<div style="margin-left:{indent_em}em; font-family:\'{cfg.font_name}\',serif; '
                f'font-size:{cfg.body_font_size_pt}pt;">{_html_escape(prefix + block.text)}</div>'
            )
            just_started_section = True
            continue

        if block.type == BlockType.BLOCKQUOTE:
            parts.append(
                f'<p style="margin-left:1.5em; font-style:italic; font-family:\'{cfg.font_name}\',serif; '
                f'font-size:{cfg.body_font_size_pt}pt;">{_html_escape(block.text)}</p>'
            )
            just_started_section = True
            continue

        if block.type == BlockType.CODE_BLOCK:
            parts.append(
                f'<p style="margin-left:0.5em; font-family:\'{cfg.monospace_font_name}\',monospace; '
                f'font-size:{max(cfg.body_font_size_pt - 1, 8)}pt;">{_html_escape(block.text)}</p>'
            )
            just_started_section = True
            continue

        if block.type == BlockType.ABSTRACT:
            parts.append('<p style="font-style:italic; margin-bottom:0;">Abstract:</p>')
            parts.append(
                f'<p style="text-align:justify; font-family:\'{cfg.font_name}\',serif; '
                f'font-size:{cfg.body_font_size_pt}pt;">{_html_escape(block.text)}</p>'
            )
            parts.append('<hr style="border-top: 1px dashed #ccc; margin: 2em 0;">')
            just_started_section = True
            continue

        # Default: body paragraph - mirrors build_document's indent rule exactly
        indent = (f"text-indent:{cfg.first_line_indent_in}in;"
                  if (cfg.enable_first_line_indent and cfg.first_line_indent_in > 0
                      and not just_started_section) else "")
        parts.append(
            f'<p style="text-align:{cfg.body_alignment}; line-height:{cfg.line_spacing}; '
            f'font-family:\'{cfg.font_name}\',serif; font-size:{cfg.body_font_size_pt}pt; {indent}">'
            f'{_html_escape(block.text)}</p>'
        )
        just_started_section = False

    return "\n".join(parts)


def build_structured_document(input_path: Path, output_path: Path,
                                cfg: StyleConfig) -> Tuple[List[str], str]:
    blocks = extract_blocks(input_path)
    if not blocks:
        raise ValueError(f"No content extracted from {input_path}")

    if cfg.detect_blockquotes:
        _detect_blockquotes(blocks)
    if cfg.detect_code_blocks:
        _detect_code_blocks(blocks)
    if cfg.detect_abstract:
        _detect_abstract(blocks)

    pipeline = HeadingPipeline(cfg)
    pipeline.run(blocks)
    _detect_list_levels(blocks)

    build_document(blocks, output_path, cfg)
    preview_html = build_preview_html(blocks, cfg)
    return pipeline.warnings, preview_html


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args():
    parser = argparse.ArgumentParser(description="Numbering-first document formatter")
    parser.add_argument("--input", "-i", required=True, help="Source .docx file")
    parser.add_argument("--output", "-o", required=True, help="Output .docx file")
    parser.add_argument("--config", "-c", default="style_config.json", help="JSON config path")
    return parser.parse_args()


def main():
    try:
        args = parse_args()
        input_path = Path(args.input).expanduser().resolve()
        output_path = Path(args.output).expanduser().resolve()
        config_path = Path(args.config).expanduser().resolve()

        print(f"Input: {input_path}", flush=True)
        print(f"Output: {output_path}", flush=True)

        if not input_path.exists():
            raise FileNotFoundError(f"Input not found: {input_path}")
        if output_path.suffix.lower() != ".docx":
            raise ValueError("Output must be a .docx file")

        cfg = _load_style_config(config_path)
        print("Running formatter...", flush=True)
        warnings, preview_html = build_structured_document(input_path, output_path, cfg)

        # Machine-readable lines the Node controller parses out of stdout.
        # Preview HTML is base64-encoded since it's multi-line and may
        # contain quotes/newlines that would break a plain stdout line.
        preview_b64 = base64.b64encode(preview_html.encode("utf-8")).decode("ascii")
        print(f"WARNINGS_JSON:{json.dumps(warnings)}", flush=True)
        print(f"PREVIEW_HTML_B64:{preview_b64}", flush=True)
        print("SUCCESS: Document formatted", flush=True)

    except Exception as e:
        print(f"ERROR: {type(e).__name__}: {e}", flush=True)
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()